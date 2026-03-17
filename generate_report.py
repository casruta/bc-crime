"""Generate a downloadable Word report from the BC-CRIME README.

Usage:
    python generate_report.py

Produces BC_Crime_Analysis_Report.docx in the BC-CRIME- directory.
Requires python-docx >= 1.1.0.
"""

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches, Pt, RGBColor, Cm, Emu

# ---------------------------------------------------------------------------
# 1. Constants
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent
README_PATH = BASE_DIR / "README.md"
OUTPUT_PATH = BASE_DIR / "BC_Crime_Analysis_Report.docx"
CHARTS_DIR = BASE_DIR / "outputs" / "charts"

COLOR_DARK_TEAL = RGBColor(0x1B, 0x3A, 0x4B)
COLOR_MEDIUM_BLUE = RGBColor(0x2C, 0x5F, 0x7C)
COLOR_GREY = RGBColor(0x60, 0x60, 0x60)
COLOR_LIGHT_GREY = RGBColor(0xF2, 0xF2, 0xF2)
COLOR_LEFT_BORDER = RGBColor(0x1B, 0x3A, 0x4B)

MAJOR_SECTIONS = {
    "Key Findings at a Glance",
    "1. Is Crime Rising?",
    "2. What Kinds of Crime Are Changing?",
    "3. How Effectively Is Crime Being Addressed?",
    "4. Where Is Crime Concentrated?",
    "5. Why Does Crime Feel Like It's Rising?",
    "Methodology",
    "Caveats",
}


# ---------------------------------------------------------------------------
# 2. Style Setup
# ---------------------------------------------------------------------------

def configure_styles(doc):
    """Configure custom styles for the Word document."""
    styles = doc.styles

    # Normal
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    # Heading 1
    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = COLOR_DARK_TEAL
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    # Heading 2
    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(16)
    h2.font.bold = True
    h2.font.color.rgb = COLOR_DARK_TEAL
    h2.paragraph_format.space_before = Pt(18)
    h2.paragraph_format.space_after = Pt(6)

    # Heading 3
    h3 = styles["Heading 3"]
    h3.font.name = "Calibri"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = COLOR_MEDIUM_BLUE
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(4)

    # Figure Caption (custom)
    if "Figure Caption" not in [s.name for s in styles]:
        caption_style = styles.add_style("Figure Caption", 1)  # 1 = paragraph
    else:
        caption_style = styles["Figure Caption"]
    caption_style.font.name = "Calibri"
    caption_style.font.size = Pt(9)
    caption_style.font.italic = True
    caption_style.font.color.rgb = COLOR_GREY
    caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_style.paragraph_format.space_after = Pt(12)

    # Key Finding (custom)
    if "Key Finding" not in [s.name for s in styles]:
        kf_style = styles.add_style("Key Finding", 1)
    else:
        kf_style = styles["Key Finding"]
    kf_style.font.name = "Calibri"
    kf_style.font.size = Pt(11)
    kf_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    kf_style.paragraph_format.left_indent = Inches(0.5)
    kf_style.paragraph_format.right_indent = Inches(0.5)
    kf_style.paragraph_format.space_after = Pt(6)
    kf_style.paragraph_format.line_spacing = 1.15

    # Code Block (custom)
    if "Code Block" not in [s.name for s in styles]:
        code_style = styles.add_style("Code Block", 1)
    else:
        code_style = styles["Code Block"]
    code_style.font.name = "Consolas"
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    code_style.paragraph_format.space_after = Pt(2)
    code_style.paragraph_format.line_spacing = 1.0


# ---------------------------------------------------------------------------
# 3. Markdown Parser
# ---------------------------------------------------------------------------

def parse_readme(lines):
    """Parse README.md lines into a list of element dicts."""
    elements = []
    state = "NORMAL"
    buffer = []
    i = 0
    n = len(lines)
    skip_toc = False
    in_details = False

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # --- Code block toggle ---
        if stripped.startswith("```"):
            if state == "CODE_BLOCK":
                elements.append({"type": "code_block", "lines": buffer})
                buffer = []
                state = "NORMAL"
                i += 1
                continue
            else:
                _flush_list(state, buffer, elements)
                buffer = []
                state = "CODE_BLOCK"
                i += 1
                continue

        if state == "CODE_BLOCK":
            buffer.append(line.rstrip())
            i += 1
            continue

        # --- Title page (lines 1, 3, 5 of the README) ---
        if i == 0 and stripped.startswith("# "):
            title = stripped[2:].strip()
            subtitle = ""
            author = ""
            j = i + 1
            while j < n:
                sl = lines[j].strip()
                if sl.startswith("**") and sl.endswith("**"):
                    subtitle = sl.strip("*").strip()
                    j += 1
                elif sl.startswith("*") and sl.endswith("*"):
                    author = sl.strip("*").strip()
                    j += 1
                elif sl == "":
                    j += 1
                elif sl == "---":
                    j += 1
                    break
                else:
                    break
            elements.append({
                "type": "title_page",
                "title": title,
                "subtitle": subtitle,
                "author": author,
            })
            i = j
            continue

        # --- Skip Table of Contents section ---
        if stripped == "## Table of Contents":
            skip_toc = True
            i += 1
            continue

        if skip_toc:
            # TOC ends at the next ## or ### List of Figures
            if stripped.startswith("### List of Figures"):
                skip_toc = False
                # fall through to handle this heading
            elif stripped.startswith("## "):
                skip_toc = False
                # fall through to handle this heading
            else:
                i += 1
                continue

        # --- <details> → appendix ---
        if stripped.startswith("<details>"):
            _flush_list(state, buffer, elements)
            buffer = []
            state = "NORMAL"
            in_details = True
            i += 1
            continue

        if stripped.startswith("<summary>"):
            text = re.sub(r"<[^>]+>", "", stripped).strip()
            if text:
                elements.append({"type": "page_break"})
                elements.append({"type": "heading", "level": 2, "text": f"Appendix: {text}"})
            i += 1
            continue

        if stripped == "</details>":
            in_details = False
            i += 1
            continue

        if stripped == "</summary>":
            i += 1
            continue

        # --- Horizontal rule ---
        if stripped == "---":
            _flush_list(state, buffer, elements)
            buffer = []
            state = "NORMAL"
            elements.append({"type": "horizontal_rule"})
            i += 1
            continue

        # --- Headings ---
        heading_match = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading_match:
            _flush_list(state, buffer, elements)
            buffer = []
            state = "NORMAL"
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 2 and text in MAJOR_SECTIONS:
                elements.append({"type": "page_break"})
            elements.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        # --- Image ---
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            _flush_list(state, buffer, elements)
            buffer = []
            state = "NORMAL"
            alt = img_match.group(1)
            path = img_match.group(2)
            elements.append({"type": "image", "alt": alt, "path": path})
            i += 1
            continue

        # --- Figure caption: *Figure N...* or *text* after an image ---
        if (stripped.startswith("*") and not stripped.startswith("**")
                and stripped.endswith("*") and not stripped.endswith("**")
                and elements and elements[-1]["type"] in ("image", "caption")):
            _flush_list(state, buffer, elements)
            buffer = []
            state = "NORMAL"
            text = stripped[1:-1].strip()
            elements.append({"type": "caption", "text": text})
            i += 1
            continue

        # --- Blockquote ---
        if stripped.startswith("> "):
            if state != "BLOCKQUOTE":
                _flush_list(state, buffer, elements)
                buffer = []
                state = "BLOCKQUOTE"
            buffer.append(stripped[2:])
            i += 1
            continue

        # --- Table ---
        if stripped.startswith("|") and stripped.endswith("|"):
            if state != "TABLE":
                _flush_list(state, buffer, elements)
                buffer = []
                state = "TABLE"
            # Skip separator rows like |---|---|
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                i += 1
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            buffer.append(cells)
            i += 1
            continue

        # --- Bullet list ---
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            if state != "BULLET_LIST":
                _flush_list(state, buffer, elements)
                buffer = []
                state = "BULLET_LIST"
            buffer.append(bullet_match.group(1))
            i += 1
            continue

        # --- Numbered list ---
        num_match = re.match(r"^\d+\.\s+(.+)$", stripped)
        if num_match:
            if state != "NUMBERED_LIST":
                _flush_list(state, buffer, elements)
                buffer = []
                state = "NUMBERED_LIST"
            buffer.append(num_match.group(1))
            i += 1
            continue

        # --- Empty line → flush ---
        if stripped == "":
            _flush_list(state, buffer, elements)
            buffer = []
            state = "NORMAL"
            i += 1
            continue

        # --- Interactive map link (plain text) ---
        link_match = re.match(
            r"^An \[interactive.*\]\(([^)]+)\)(.*)$", stripped
        )
        if link_match:
            _flush_list(state, buffer, elements)
            buffer = []
            state = "NORMAL"
            filename = Path(link_match.group(1)).name
            rest = link_match.group(2).strip()
            text = f"An interactive version of the jurisdiction data is available (see {filename}){rest}"
            elements.append({"type": "paragraph", "text": text})
            i += 1
            continue

        # --- Regular paragraph ---
        _flush_list(state, buffer, elements)
        buffer = []
        state = "NORMAL"
        elements.append({"type": "paragraph", "text": stripped})
        i += 1

    _flush_list(state, buffer, elements)
    return elements


def _flush_list(state, buffer, elements):
    """Flush any accumulated list/table/blockquote buffer into elements."""
    if not buffer:
        return
    if state == "TABLE":
        elements.append({"type": "table", "rows": list(buffer)})
    elif state == "BLOCKQUOTE":
        elements.append({"type": "blockquote", "lines": list(buffer)})
    elif state == "BULLET_LIST":
        elements.append({"type": "bullet_list", "items": list(buffer)})
    elif state == "NUMBERED_LIST":
        elements.append({"type": "numbered_list", "items": list(buffer)})
    buffer.clear()


# ---------------------------------------------------------------------------
# 4. Inline Formatting
# ---------------------------------------------------------------------------

_INLINE_PATTERN = re.compile(
    r"(\*\*(.+?)\*\*)"         # bold
    r"|(\*(.+?)\*)"            # italic
    r"|(`([^`]+)`)"            # inline code
    r"|(\[([^\]]+)\]\([^)]+\))" # link (render text only)
)


def _add_formatted_runs(paragraph, text):
    """Split text into bold, italic, code, and plain runs."""
    last = 0
    for m in _INLINE_PATTERN.finditer(text):
        start = m.start()
        if start > last:
            paragraph.add_run(text[last:start])

        if m.group(2) is not None:
            # bold
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(4) is not None:
            # italic
            run = paragraph.add_run(m.group(4))
            run.italic = True
        elif m.group(6) is not None:
            # inline code
            run = paragraph.add_run(m.group(6))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif m.group(8) is not None:
            # link text
            paragraph.add_run(m.group(8))

        last = m.end()

    if last < len(text):
        paragraph.add_run(text[last:])


def _add_left_border(paragraph):
    """Add a colored left border to a paragraph via OxmlElement."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4")
    left.set(qn("w:color"), "1B3A4B")
    pBdr.append(left)
    pPr.append(pBdr)


def _set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
# 5. Document Builder
# ---------------------------------------------------------------------------

def build_document(elements):
    """Build a Word document from parsed elements."""
    doc = Document()
    configure_styles(doc)

    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("BC Crime Analysis Report")
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GREY
    run.font.name = "Calibri"

    # Footer with page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_page_number(fp)

    stats = {"paragraphs": 0, "tables": 0, "images": 0, "images_missing": 0}

    for elem in elements:
        etype = elem["type"]

        if etype == "title_page":
            _build_title_page(doc, elem)
            stats["paragraphs"] += 3

        elif etype == "page_break":
            doc.add_page_break()

        elif etype == "heading":
            level = elem["level"]
            style = f"Heading {level}"
            p = doc.add_heading(elem["text"], level=level)
            # Re-apply color since add_heading resets
            for run in p.runs:
                if level <= 2:
                    run.font.color.rgb = COLOR_DARK_TEAL
                else:
                    run.font.color.rgb = COLOR_MEDIUM_BLUE
            stats["paragraphs"] += 1

        elif etype == "image":
            img_path = BASE_DIR / elem["path"]
            if img_path.exists():
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(str(img_path), width=Inches(6.0))
                stats["images"] += 1
            else:
                p = doc.add_paragraph(f"[Image not found: {elem['path']}]")
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                stats["images_missing"] += 1
                print(f"  WARNING: Missing image: {img_path}", file=sys.stderr)

        elif etype == "caption":
            p = doc.add_paragraph(style="Figure Caption")
            _add_formatted_runs(p, elem["text"])
            stats["paragraphs"] += 1

        elif etype == "table":
            _build_table(doc, elem["rows"])
            stats["tables"] += 1

        elif etype == "blockquote":
            for line in elem["lines"]:
                p = doc.add_paragraph(style="Key Finding")
                _add_formatted_runs(p, line)
                _add_left_border(p)
            stats["paragraphs"] += len(elem["lines"])

        elif etype == "bullet_list":
            for item in elem["items"]:
                p = doc.add_paragraph(style="List Bullet")
                _add_formatted_runs(p, item)
            stats["paragraphs"] += len(elem["items"])

        elif etype == "numbered_list":
            for item in elem["items"]:
                p = doc.add_paragraph(style="List Number")
                _add_formatted_runs(p, item)
            stats["paragraphs"] += len(elem["items"])

        elif etype == "code_block":
            for code_line in elem["lines"]:
                p = doc.add_paragraph(style="Code Block")
                run = p.add_run(code_line)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                # Light grey background on each code line
                _set_paragraph_shading(p, "F2F2F2")
            stats["paragraphs"] += len(elem["lines"])

        elif etype == "horizontal_rule":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)

        elif etype == "paragraph":
            p = doc.add_paragraph()
            _add_formatted_runs(p, elem["text"])
            stats["paragraphs"] += 1

    return doc, stats


def _build_title_page(doc, elem):
    """Build a centered title page with title, subtitle, and author."""
    # Add some vertical spacing
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(elem["title"])
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = COLOR_DARK_TEAL
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(18)

    # Subtitle
    if elem["subtitle"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(elem["subtitle"])
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_GREY
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(12)

    # Author
    if elem["author"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(elem["author"])
        run.font.size = Pt(12)
        run.font.italic = True
        run.font.color.rgb = COLOR_GREY
        run.font.name = "Calibri"
        p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()


def _build_table(doc, rows):
    """Build a formatted Word table from parsed rows."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Try to apply a built-in table style
    try:
        table.style = "Light Grid Accent 1"
    except KeyError:
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < n_cols:
                cell = row.cells[c_idx]
                p = cell.paragraphs[0]
                _add_formatted_runs(p, cell_text)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                # Bold first row (header)
                if r_idx == 0:
                    for run in p.runs:
                        run.bold = True
                    _set_cell_shading(cell, "1B3A4B")
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Allow header row to repeat across pages
    if len(rows) > 1:
        _set_repeat_header_row(table)


def _set_repeat_header_row(table):
    """Mark the first row as a header row that repeats across pages."""
    tr = table.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _set_paragraph_shading(paragraph, color_hex):
    """Set background shading on a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    pPr.append(shading)


def _add_page_number(paragraph):
    """Add a centered page number field to a footer paragraph."""
    run = paragraph.add_run()
    run.font.size = Pt(8)
    run.font.color.rgb = COLOR_GREY
    run.font.name = "Calibri"

    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    run._r.append(fldChar1)

    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = " PAGE "
    run._r.append(instrText)

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar2)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print("Reading README.md...")
    lines = README_PATH.read_text(encoding="utf-8").splitlines()

    print(f"Parsing {len(lines)} lines...")
    elements = parse_readme(lines)

    print("Building Word document...")
    doc, stats = build_document(elements)

    print(f"Saving to {OUTPUT_PATH}...")
    doc.save(str(OUTPUT_PATH))

    print("\nReport generated successfully!")
    print(f"  Paragraphs: {stats['paragraphs']}")
    print(f"  Tables:     {stats['tables']}")
    print(f"  Images:     {stats['images']}")
    if stats["images_missing"]:
        print(f"  Missing:    {stats['images_missing']}")
    print(f"\n  Output: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
